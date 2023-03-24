import pickle
from pathlib import Path 
import streamlit_authenticator as stauth
import streamlit as st 
import docx
import tempfile
import os 
import openai

# Set OpenAI API key
openai.api_key = os.environ["OPENAI_API_KEY"]

names = ['Edwar Benavente', 'Eduardo Benavente', 'Adela Ingiltupa']
usernames = ['master','edube', 'adela']


file_path = Path(__file__).parent / "hashed_pw.pkl"

with file_path.open("rb") as file:
    hashed_passwords = pickle.load(file)

authenticator = stauth.Authenticate(names, usernames, hashed_passwords,"autenticador", "abcdef", cookie_expiry_days=30)

name, authentication_status, username = authenticator.login("Login", "main")

if authentication_status == False:
    st.error('Usuario/Contraseña es incorrecto')

if authentication_status == None:
    st.warning("Por favor introducir su usuario y contraseña")

if authentication_status:
    
    import pandas as pd
    import streamlit as st
    import openai

    openai.api_key = os.environ["OPENAI_API_KEY"]

    def generate_learning_plan(grade_level, subject_area, topic):
        # Use ChatGPT to generate the learning plan based on user inputs
        
        prompts = [f"Eres un profesor de educación primaria de {grade_level}. Diseñarás una sesión de aprendizaje detallada del curso {subject_area} sobre el tema {topic}. Debes detallar incorporar estos componentes: Propósitos de aprendizaje, Preparación de la actividad, Desarrollo de la estrategia, Estrategia de Evaluación (Rúbrica) y Cierre (Reflexión).",    "En cuanto al propósito de aprendizaje, deberás incluir como subtítulos y explicar: competencias y capacidades que trabajarás, los desempeños esperados, los criterios de evaluación, las evidencias e instrumentos de evaluación que utilizarás para medir el éxito de la sesión.",    "En cuanto a la preparación de la actividad, debes incluir como subtítulos y explicar: las tareas y actividades que los estudiantes realizarán antes, durante y después de la sesión, y cómo estas se alinean con los objetivos de aprendizaje. Además, brindarás una lista de recursos y materiales que se utilizarán en esta sesión.",    "En cuanto al desarrollo de la estrategia, debes incluir como subtítulos y explicar: cómo motivarás a tus estudiantes, cómo integrarás sus saberes previos, cómo presentarás la situación problemática y cómo desarrollarás el tema. Es importante en esta sección detallar la situación problemática, la cual deberá formularse en forma de pregunta.",    "Por último, en cuanto al cierre, deberás describir cómo ayudarás a tus estudiantes a reflexionar sobre lo que han aprendido y cómo fomentarás la metacognición. Como consideraciones finales, incluye las rúbricas de evaluación y tablas donde consideres relevante. Todos los componentes y subcomponentes deben estar escritos y fundamentados detalladamente y debe cumplir con los requisitos. El resultado que me darás debe ser menor a 1700 palabras."]

        responses = []
        for p in prompts:
            response = openai.Completion.create(
                engine="text-davinci-003", prompt=p, max_tokens=3900, n=1, stop=None, temperature=0.5,
            )
            learning_plan = response.choices[0].text
            responses.append(learning_plan)

        learning_plan = "".join(responses)
        
        return learning_plan  

    # Create the Streamlit user interface
    st.title("Generar una Nueva Sesión de Aprendizaje")
    # Replace <phone_number> with your phone number in international format
    whatsapp_url = "https://wa.me/447517438058"
    
    # Display the WhatsApp link and personal WhatsApp
    st.write(f"¿Necesitas ayuda? Contáctanos en [WhatsApp]({whatsapp_url}). Te ayudaremos de inmediato!")
    
    
    grade_level = st.selectbox("Seleccione el grado: ", ["1er grado", "2do grado", "3er grado", "4to grado", "5to grado", "6to grado"])
    subject_area = st.selectbox("Seleccione el curso: ", ['Comunicación', 'Matemática', 'Personal Social', 'Ciencia y Tecnología', 'Arte', 'Educación Religiosa', 'Educación Física', 'Tutoría'])
    duration = st.slider("Seleccione la duración de la sesión (en minutos)", min_value=30, max_value=120, value=60)
    topic = st.text_input("Ingrese el tema de la sesión de aprendizaje, debe ser lo más específico posible (Ej: La célula animal, Ecuaciones de primer grado, Operaciones con conjuntos, La cultura Chavín Peruana):")
    st.warning("¡Presione el botón y en 3 minutos tendrá su Sesión de Aprendizaje!")
    # Generate the learning plan based on user inputs
    if st.button("Generar Sesión de Aprendizaje"):
        progress_bar = st.progress(0)
        learning_plan = generate_learning_plan(grade_level, subject_area, topic)
        progress_bar.progress(50)
        st.write(learning_plan)
        progress_bar.progress(100)
        
        # Convert the response to a Word document
        doc = docx.Document()
        doc.add_paragraph(learning_plan)

        # Convert the document to binary data
        with tempfile.NamedTemporaryFile(delete=False) as f:
            doc.save(f.name)
            f.flush()
            data = f.read()

        # Download the document
        st.download_button(
            label="Descargar respuesta como documento de Word",
            data=data,
            file_name="respuesta.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

   
        # Add text input field to the sidebar
with st.sidebar:
    st.subheader('¿Necesitas ideas para tu sesión?')
    st.write('Aquí puedes preguntar lo que quieras ¡Debe ser detallado! Ejemplos:')
    st.write('  * Dame 5 problemas de sumas y restas para 2do de primaria con contexto.')
    st.write('  * Texto del Día del Agua para 5to de primaria en formato de cuento, con preguntas de comprensión lectora')  
    user_input = st.sidebar.text_input("En qué te puedo ayudar:")

    # Add button to the sidebar
    if st.sidebar.button("Preguntar"):
        # Call the OpenAI API
        response = openai.Completion.create(
            engine="text-davinci-003",
            prompt=user_input,
            max_tokens=1024,
            n=1,
            stop=None,
            temperature=0.3,
        )

        # Convert the response to a Word document
        bot_response = response.choices[0].text
        doc = docx.Document()
        doc.add_paragraph(bot_response)

        # Convert the document to binary data
        with tempfile.NamedTemporaryFile(delete=False) as f:
            doc.save(f.name)
            f.flush()
            data = f.read()

        # Download the document
        st.download_button(
            label="Descargar respuesta como documento de Word",
            data=data,
            file_name="respuesta.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        # Display the bot response text in the sidebar
        st.sidebar.text(bot_response)
        